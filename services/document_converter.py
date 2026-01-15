import os
import tempfile
import subprocess
from pathlib import Path
from typing import Optional, Tuple
import logging
from docx import Document
import shutil
import email
from email import policy
from email.parser import BytesParser

logger = logging.getLogger("document_ai")

class DocumentConverter:
    """Service for converting documents to formats supported by Document AI"""
    
    SUPPORTED_IMAGE_FORMATS = {'.jpg', '.jpeg', '.png', '.gif', '.tiff', '.tif', '.bmp', '.webp'}
    SUPPORTED_DOCUMENT_FORMATS = {'.pdf'}
    # Exhaustive list of formats we can try to convert using LibreOffice or internal parsers
    CONVERTIBLE_FORMATS = {
        '.docx', '.doc', '.rtf', '.odt', '.txt', '.wps', '.wpd',  # Documents
        '.dot', '.dotx', '.docm', '.dox',                         # Templates/Macros/Legacy
        '.pptx', '.ppt', '.odp',                                  # Presentations
        '.xlsx', '.xls', '.csv', '.ods',                          # Spreadsheets
        '.eml', '.msg',                                           # Emails
        '.html', '.htm', '.xml'                                   # Web/structured
    }
    
    @classmethod
    def get_supported_formats(cls) -> set:
        """Get all supported and convertible file formats"""
        return cls.SUPPORTED_IMAGE_FORMATS | cls.SUPPORTED_DOCUMENT_FORMATS | cls.CONVERTIBLE_FORMATS
    
    @classmethod
    def is_supported_format(cls, file_path: str) -> bool:
        """Check if file format is supported or convertible"""
        # If the file has no extension or we want to be permissive, we can return True
        # But for now, let's rely on the expanded list.
        file_ext = Path(file_path).suffix.lower()
        return file_ext in cls.get_supported_formats()
    
    @classmethod
    def needs_conversion(cls, file_path: str) -> bool:
        """Check if file needs conversion to be processed by Document AI"""
        file_ext = Path(file_path).suffix.lower()
        return file_ext in cls.CONVERTIBLE_FORMATS

    @classmethod
    def convert_email_to_pdf(cls, email_path: str, output_dir: Optional[str] = None) -> str:
        """
        Convert .eml file to PDF by extracting body (HTML/Text) and converting via LibreOffice.
        Note: .msg support requires additional libraries, so for now we treat .msg same as .eml 
        (which might fail if it's binary OLE) or assume it's text-based.
        """
        try:
            logger.info(f"📧 Converting Email to PDF: {email_path}")
            
            if output_dir is None:
                output_dir = tempfile.gettempdir()
            
            # Parse email content
            with open(email_path, 'rb') as fp:
                msg = BytesParser(policy=policy.default).parse(fp)

            # Get body content
            body_content = ""
            body = msg.get_body(preferencelist=('html', 'plain'))
            
            if body:
                body_content = body.get_content()
                content_type = body.get_content_type() 
            else:
                # Fallback if get_body fails
                if msg.is_multipart():
                    for part in msg.walk():
                        if part.get_content_type() == "text/html":
                            body_content = part.get_content()
                            content_type = "text/html"
                            break
                        elif part.get_content_type() == "text/plain":
                            body_content = part.get_content()
                            content_type = "text/plain"
                else:
                    body_content = msg.get_content()
                    content_type = msg.get_content_type()

            # Create an intermediate HTML file
            # If plain text, wrap in <pre> or basic HTML
            if "html" not in content_type:
                html_content = f"<html><body><pre>{body_content}</pre></body></html>"
            else:
                html_content = body_content
                
            # Add header info to the top of the HTML (Subject, From, Date)
            headers_html = "<div style='border-bottom: 1px solid #ccc; margin-bottom: 20px; padding-bottom: 10px;'>"
            if msg['subject']: headers_html += f"<b>Subject:</b> {msg['subject']}<br>"
            if msg['from']: headers_html += f"<b>From:</b> {msg['from']}<br>"
            if msg['to']: headers_html += f"<b>To:</b> {msg['to']}<br>"
            if msg['date']: headers_html += f"<b>Date:</b> {msg['date']}<br>"
            headers_html += "</div>"
            
            final_html = headers_html + html_content
            
            # Save intermediate HTML
            stem = Path(email_path).stem
            html_path = os.path.join(output_dir, f"{stem}_temp.html")
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(final_html)
                
            # Convert HTML to PDF using LibreOffice
            # We can re-use convert_office_to_pdf logic or call it directly
            # Since convert_docx_to_pdf is being renamed/generalized, let's call the generic converter
            # But here we invoke LibreOffice directly on the HTML file
            
            return cls.convert_office_to_pdf(html_path, output_dir)
            
        except Exception as e:
            logger.error(f"❌ Error converting email to PDF: {str(e)}")
            raise Exception(f"Email conversion failed: {str(e)}")

    @classmethod
    def convert_office_to_pdf(cls, file_path: str, output_dir: Optional[str] = None) -> str:
        """
        Generic converter using LibreOffice (Headless).
        Supports DOCX, DOC, RTF, ODT, HTML, TXT, XLSX, PPTX, etc.
        """
        try:
            logger.info(f"🔄 Converting file to PDF using LibreOffice: {file_path}")
            
            # Validate input file
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Input file not found: {file_path}")
            
            # Determine output directory
            if output_dir is None:
                output_dir = tempfile.gettempdir()
            
            # Ensure output directory exists
            os.makedirs(output_dir, exist_ok=True)
            
            # Use LibreOffice to convert to PDF
            # LibreOffice handles many formats automatically
            cmd = [
                "libreoffice", 
                "--headless",           # Run without GUI
                "--convert-to", "pdf",  # Convert to PDF format
                "--outdir", output_dir, # Output directory
                file_path              # Input file
            ]
            
            logger.info(f"🚀 Running LibreOffice command: {' '.join(cmd)}")
            
            # Run the conversion
            result = subprocess.run(
                cmd, 
                capture_output=True, 
                text=True, 
                timeout=60  # 60 second timeout
            )
            
            if result.returncode != 0:
                error_msg = f"LibreOffice conversion failed with return code {result.returncode}"
                if result.stderr:
                    error_msg += f": {result.stderr}"
                logger.error(f"❌ {error_msg}")
                raise Exception(error_msg)
            
            # Determine expected PDF path
            input_stem = Path(file_path).stem
            pdf_path = os.path.join(output_dir, f"{input_stem}.pdf")
            
            # Verify PDF was created
            if not os.path.exists(pdf_path):
                raise Exception(f"PDF conversion failed - expected output file not found: {pdf_path}")
            
            pdf_size = os.path.getsize(pdf_path)
            if pdf_size == 0:
                raise Exception("PDF conversion failed - output file is empty")
            
            logger.info(f"✅ Successfully converted to PDF: {pdf_path}")
            logger.info(f"📏 PDF file size: {pdf_size} bytes")
            
            # Log LibreOffice output if available
            if result.stdout:
                logger.info(f"📋 LibreOffice output: {result.stdout.strip()}")
            
            return pdf_path
            
        except subprocess.TimeoutExpired:
            logger.error("❌ LibreOffice conversion timed out after 60 seconds")
            raise Exception("File to PDF conversion timed out")
        except Exception as e:
            logger.error(f"❌ Error converting file to PDF: {str(e)}")
            raise Exception(f"File to PDF conversion failed: {str(e)}")
    
    @classmethod
    def extract_text_from_docx(cls, docx_path: str) -> str:
        """
        Extract plain text from DOCX file as fallback option
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            logger.info(f"📄 Extracting text from DOCX: {docx_path}")
            
            document = Document(docx_path)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            extracted_text = "\n".join(text_content)
            logger.info(f"✅ Extracted {len(extracted_text)} characters from DOCX")
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"❌ Error extracting text from DOCX: {str(e)}")
            raise Exception(f"DOCX text extraction failed: {str(e)}")
    
    @classmethod
    def convert_document(cls, input_path: str, target_format: str = "pdf") -> Tuple[str, bool]:
        """
        Convert document to target format if needed
        
        Args:
            input_path: Path to input document
            target_format: Target format ('pdf' or 'text')
            
        Returns:
            Tuple of (converted_file_path, was_converted)
        """
        file_ext = Path(input_path).suffix.lower()
        
        # If already in supported format, return as-is
        if file_ext in (cls.SUPPORTED_IMAGE_FORMATS | cls.SUPPORTED_DOCUMENT_FORMATS):
            logger.info(f"✅ File already in supported format: {file_ext}")
            return input_path, False
        
        # Convert based on file type and target format
        # Use our exhaustive CONVERTIBLE_FORMATS set for checking
        if file_ext in cls.CONVERTIBLE_FORMATS:
            if target_format == "pdf":
                if file_ext in ['.eml', '.msg']:
                    # Special handling for emails
                    converted_path = cls.convert_email_to_pdf(input_path)
                else:
                    # General office processing (Word, Excel, PPT, Text, HTML...)
                    converted_path = cls.convert_office_to_pdf(input_path)
                return converted_path, True
            elif target_format == "text":
                # For text extraction, we return original path as fallback logic usually handles text extraction
                return input_path, False
        
        raise ValueError(f"Unsupported file format for conversion: {file_ext}")
    
    @classmethod
    def cleanup_converted_file(cls, file_path: str, was_converted: bool):
        """Clean up converted file if it was created during conversion"""
        if was_converted and os.path.exists(file_path):
            try:
                os.unlink(file_path)
                logger.info(f"🗑️ Cleaned up converted file: {file_path}")
            except Exception as e:
                logger.warning(f"⚠️ Could not clean up converted file {file_path}: {str(e)}")
    
    def convert_to_pdf(self, content: bytes, filename: str) -> bytes:
        """
        Convert document bytes to PDF bytes for preview.
        
        Args:
            content: Binary content of the input file
            filename: Original filename (with extension) to determine format
            
        Returns:
            PDF content as bytes
        """
        temp_dir = tempfile.mkdtemp()
        input_filename = Path(filename).name
        try:
            input_path = os.path.join(temp_dir, input_filename)
            with open(input_path, 'wb') as f:
                f.write(content)
            
            logger.info(f"🔄 Converting {input_filename} to PDF using temporary path: {input_path}")
            
            converted_path, was_converted = DocumentConverter.convert_document(input_path, "pdf")
            
            with open(converted_path, 'rb') as f:
                pdf_bytes = f.read()
            
            if not pdf_bytes:
                raise Exception("PDF conversion resulted in empty content")
            
            logger.info(f"✅ Converted to PDF: {len(pdf_bytes)} bytes")
            return pdf_bytes
            
        except Exception as e:
            logger.error(f"❌ Error in convert_to_pdf for {filename}: {str(e)}")
            raise
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)